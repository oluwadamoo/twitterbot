
import snscrape.modules.twitter as sntwitter
import docx
from datetime import datetime


# Created a list to append all tweet attributes(data)
# attributes_container = []
my_doc = docx.Document()

spool_type = input("What parameter do you want to spool by?\n (Enter 1 for 'username' and Enter 2 for 'Topic'): ")
username = ''
topic = ''
search_query = ''

if spool_type == '1':
    username = input("Enter the username: ")
    search_query = 'from:{}'.format(username)
elif spool_type == '2':
    topic = input("Enter topic you want to search: ")
    search_query = topic
else:
    print("You never get time")

no_of_tweets = int(input('Enter the number of tweets you want to spool: '))

if no_of_tweets:
    no_of_tweets = no_of_tweets
else:
    no_of_tweets = 1000

date_from = input("Enter start date in the format 2021-07-05: ")
date_to = input("Enter end date in the format 2021-07-05: ")
dt = datetime.now()
ts = datetime.timestamp(dt)

# Using TwitterSearchScraper to scrape data and append tweets to list

if search_query:
    search = '{} since:{} until:{}'.format(search_query, date_from, date_to)
    for i, tweet in enumerate(sntwitter.TwitterSearchScraper(search).get_items()):
        if i > no_of_tweets:
            break
        if '@' in tweet.rawContent:
            continue
        my_doc.add_paragraph(tweet.user.username)
        my_doc.add_paragraph(str(tweet.date))
        my_doc.add_paragraph(tweet.rawContent)
        my_doc.add_paragraph("\n")
        # attributes_container.append(tweet.rawContent)
else:
    print('You never get time')
my_doc.save("results-{}.docx".format(int(ts)))
